# encoding: utf-8

from PIL import ImageGrab, Image
import numpy as np
import mouse
from time import sleep
import pytesseract
import PySimpleGUI as sg
from winsound import Beep
from os import makedirs, path, remove, environ, listdir
from datetime import datetime
import webbrowser
from configparser import ConfigParser
from docx import Document
from docx.shared import Cm
from random import choice
import sounddevice as sd
import soundfile as sf
import threading
import speech_recognition as sr
import wave



#img = Image.fromarray(screen, 'RGB')
#img.save('my.png')
#img.show()

#import os # MAC I LINUX
#duration = 1  # seconds
#freq = 440  # Hz
#os.system('play -nq -t alsa synth {} sine {}'.format(duration, freq))

#enhancer1 = ImageEnhance.Sharpness(img)
#enhancer2 = ImageEnhance.Contrast(img)
#img_edit = enhancer1.enhance(20.0)
#img_edit = enhancer2.enhance(1.5)



#Write the above sections to config.ini file
#with open('config.ini', 'w') as conf:
#    config_object.write(conf)


# WCZYTAJ CONFIGI #

config_object = ConfigParser()
config_object.read("config.ini")

okno_config = config_object["okno"]
tesseract_config = config_object["tesseract"]
preferencje_config = config_object["preferencje"]
audio_config = config_object["audio"]




# SETUP #

def rec_unlimited(kill_event, czas):
    #edited from https://github.com/spatialaudio/python-sounddevice/blob/master/examples/rec_unlimited.py
    """Create a recording with arbitrary duration.
    The soundfile module (https://PySoundFile.readthedocs.io/) has to be installed!
    """
    import queue
    q = queue.Queue()

    def callback(indata, frames, time, status):
        """This is called (from a separate thread) for each audio block."""
        if status:
            print(status, file=sys.stderr)
        q.put(indata.copy())
        
        
        # Make sure the file is opened before recording anything:
    with sf.SoundFile("temp_"+czas+".wav", mode='x', samplerate=48000,
                      channels=2, subtype='PCM_16') as file: #sf.default_subtype('WAV')
        with sd.InputStream(samplerate=48000, device=2, #TUTAJ TUTAJ TUTAJ TUTAJ TUTAJ#
                            channels=2, callback=callback):
            print('nagrywanie_audio_on')
            while not kill_event.is_set():
                file.write(q.get())


def rec_szum(czas, recognizer):
    profil_szumu = sd.rec(int(czas * 48000), samplerate=48000, device=2, channels=2)
    sf.write('temp_szum.wav', profil_szumu, 48000)
    with sr.AudioFile("temp_szum.wav") as source:
        recognizer.adjust_for_ambient_noise(source, duration=(czas-0.1))


def tekst_mowy(czas_pliku, lang_a):
    with open("temp_"+czas_pliku+".wav", 'rb') as pcmfile:
        pcmdata = pcmfile.read()
    # naprawianie .wav #
    with wave.open("temp_"+czas_pliku+".wav", 'wb') as wavfile:
        wavfile.setparams((2, 2, 48000, 1, 'NONE', 'NONE'))
        wavfile.writeframes(pcmdata)
    with sr.AudioFile("temp_"+czas_pliku+".wav") as source:
        recorded_audio = recognizer.listen(source)
#        recognizer.adjust_for_ambient_noise(source, duration=0.1)
        try:
            return recognizer.recognize_google(recorded_audio, language=lang_a)
        except sr.UnknownValueError as uve:
            print(uve)
            return "### Nie rozpoznano mowy. ###"


def dodaj_transkrypt(dokument, czas_pliku, lang_a):
    dokument.add_paragraph(czas_pliku +' >>> '+ tekst_mowy(czas_pliku, lang_a), style='Quote')
    

def wyznacz_okno():
    print('a:', end=' ')
    while True:
        if mouse.is_pressed("left"):
            ax, ay = mouse.get_position()
            Beep(1120, 90)
            break
    print(ax, ay)
    sleep(0.3)
    print('b:', end=' ')
    while True:
        if mouse.is_pressed("left"):
            bx, by = mouse.get_position()
            Beep(1450, 120)
            break
    print(bx, by)
    return ax, ay, bx, by


def teraz():
    return datetime.now().strftime("%m,%e-%H,%M,%S")


##############################################################

#fs = 48000  # Sample rate
#seconds = 3  # Duration of recording
#sd.default.device = 2  # Speakers full name here
#
#myrecording = sd.rec(int(seconds * fs), samplerate=fs, channels=2)
#sd.wait()  # Wait until recording is finished
#write('output.wav', fs, myrecording)  # Save as WAV file 


#rec_unlimited(threading.Event(), "123,123,")
#exit()



# wczytaj z config #

ax, ay, bx, by = int(okno_config["ax"]), int(okno_config["ay"]), int(okno_config["bx"]), int(okno_config["by"])
pytesseract.pytesseract.tesseract_cmd = tesseract_config["exe path"] #r"D:\PROG_\Tesseract\tesseract.exe"
langs = tesseract_config["langs"]
lang_a = audio_config["lang"]
prog_slajdu = float(okno_config["progowanie"])
if (czy_docx := preferencje_config["output"] == 'docx'):
    dokument = Document()
if (czy_audio := audio_config["transkrypcja"] =='1'):
    recognizer = sr.Recognizer()
czy_autozapis = preferencje_config["autozapis"] == '1'
czy_beep = preferencje_config["bip bop"] == '1'
szum_czas = float(audio_config["czas_szumu"])


environ["TESSDATA_PREFIX"] = tesseract_config["tessdata path"]

# scieżka na notatki #
newpath = r'.\auto_notatki' 
if not path.exists(newpath):
    makedirs(newpath)
    



# THEME I LAYOUT #

# themes #
sg.LOOK_AND_FEEL_TABLE['Szymon'] = {'BACKGROUND': '#5b1b35', 
                                        'TEXT': '#cecedd', 
                                        'INPUT': '#339966', 
                                        'TEXT_INPUT': '#000000', 
                                        'SCROLL': '#99CC99', 
                                        'BUTTON': ('#e0d9dc', '#4c4247'), 
                                        'PROGRESS': ('#D1826B', '#CC8019'), 
                                        'BORDER': 1, 'SLIDER_DEPTH': 0, 'PROGRESS_DEPTH': 0, }

sg.LOOK_AND_FEEL_TABLE['Ola'] = {'BACKGROUND': '#d2dbe0', 
                                        'TEXT': '#1c2328', 
                                        'INPUT': '#339966', 
                                        'TEXT_INPUT': '#000000', 
                                        'SCROLL': '#99CC99', 
                                        'BUTTON': ('#1c2328', '#9ea8af'), 
                                        'PROGRESS': ('#D1826B', '#CC8019'), 
                                        'BORDER': 1, 'SLIDER_DEPTH': 0, 'PROGRESS_DEPTH': 0, }


hellows = ["Witaj zbłąkana istoto", "ZOBACZ JAK", "Odkrył jeden prosty sposób", "Henlo",
           "Serwus", "Zupa nie ma definicji szczerze"]

sg.theme(preferencje_config["motyw"])


layout = [[sg.Text("status:"), sg.Text(choice(hellows), key='stat')],
          [sg.Button("Wyznacz okno", key='okno'), sg.Button("Notuj", key='onoff', button_color=('white','crimson')), sg.Button("Instrukcja"), sg.Button("Wyjdź")]]
if czy_audio:
    layout[1].insert(1, sg.Button("Profiluj szum", key='szum'))


window = sg.Window("Notuj za rynkunsa", layout, keep_on_top = True,)




# EVENT LOOP #

while True:
    event, values = window.read()
    
    
    if event == "okno":
        window['stat'].update('Zaznacz narożniki')
        _ = window.read(timeout=3)
        ax, ay, bx, by = wyznacz_okno()
        # zapisz zmiany #
        okno_config["ax"], okno_config["ay"], okno_config["bx"], okno_config["by"] = str(ax), str(ay), str(bx), str(by)
        with open('config.ini', 'w') as conf:
            config_object.write(conf)
        window['stat'].update('Zarejestrowano okno')
    
    
    if event == "Instrukcja":
        webbrowser.open("README.txt")
        
    if event == "szum":
        print(szum_czas, 'sek profilowania szumu')
        window['stat'].update('Nagrywam profil szumu...')
        rec_szum(szum_czas, recognizer)
        window['stat'].update('Szum sprofilowany!')
    
    
    if event == "onoff":
        print("___ ___ ___")
        czas = teraz()
        if not czy_autozapis:
            window['okno'].update('Zanotuj okno')
        with open(newpath + '\\notka-' + czas + '.txt', 'w', encoding="utf-8") as plik:            
            if czy_docx:
                dokument.add_paragraph("Automatyczna notatka z "+czas+"\n\n")
            plik.write("Automatyczna notatka z "+czas+"\n\n")
            window['stat'].update('Notuję...')
            window['onoff'].update('Zakończ')
            zrzut = np.array(ImageGrab.grab(bbox=(ax,ay, bx,by)))
            img = Image.fromarray(zrzut, 'RGB')
            scrape = str((pytesseract.image_to_string(img, lang=langs)))
            if czy_docx:
                img = Image.fromarray(zrzut, 'RGB')
                img.save('temp_notuj.png')
                dokument.add_picture('temp_notuj.png', width=Cm(16))
                dokument.add_paragraph(scrape.strip())
            plik.write(scrape) # SAVING
            old_scrape = scrape
            
            # audio #
            if czy_audio:
                killswitch = threading.Event()
                t = threading.Thread(target=rec_unlimited, args=(killswitch,c:=teraz()))
                t.start()
                window['szum'].update(visible=False)
                
            while True:
                event, values = window.read(timeout=250)
                
                old_zrzut = zrzut
                zrzut = np.array(ImageGrab.grab(bbox=(ax,ay, bx,by)))
                #zmiana = np.mean((zrzut-old_zrzut)**2)
                window['stat'].update('Notuję... *'+str(round(zmiana:=np.mean((zrzut-old_zrzut)**2), 2)))
                if zmiana > prog_slajdu: # COMPARE 
                    img = Image.fromarray(zrzut, 'RGB')
                    scrape = str((pytesseract.image_to_string(img, lang=langs)))
                    if scrape != old_scrape:
                        if czy_autozapis: # SAVING
                            if czy_audio:
                                killswitch.set()
                                #t.join()
                                t2 = threading.Thread(target=dodaj_transkrypt, args=(dokument, c, lang_a))
                                
                                killswitch = threading.Event()
                                t = threading.Thread(target=rec_unlimited, args=(killswitch, c:=teraz()))
                                t.start()
                            if czy_docx:
                                img = Image.fromarray(zrzut, 'RGB')
                                img.save('temp_notuj.png')
                                dokument.add_picture('temp_notuj.png', width=Cm(16))
                                dokument.add_paragraph(scrape.strip())
                                t2.start()
                            plik.write(scrape)
                            window['stat'].update('Zanotowano... *'+str(round(zmiana, 2)))
                            if czy_beep: Beep(3000, 50)
                        old_scrape = scrape
                
                if event == "onoff" or event == "Wyjdź":
                    if czy_audio:
                        killswitch.set()
                        #t.join()
                        #analiza#
#                        plik.write('<<<'+trsk+'>>>')
                        window['szum'].update(visible=True)
                    if czy_docx:
                        dodaj_transkrypt(dokument, c, lang_a)
                        dokument.save(newpath + '\\notka-' + czas + '.docx')
                    # cleanup #
                    for fnazwa in listdir('.'):
                        if fnazwa.startswith("temp_"):
                            remove(path.join(fnazwa))
                            
                    print('done.')
                    window['stat'].update('Dość.')
                    if czy_beep:
                        Beep(2200, 90)
                        Beep(2600, 90)
                    window['onoff'].update('Notuj')
                    window['okno'].update('Wyznacz okno')
                    break
                
                if event == sg.WIN_CLOSED:
                    break
                
                if event == "okno":
                    if not czy_autozapis:
                        if czy_audio:
                            killswitch.set()
                            #t.join()
                            t2 = threading.Thread(target=dodaj_transkrypt, args=(dokument, c, lang_a))
                            
                            killswitch = threading.Event()
                            t = threading.Thread(target=rec_unlimited, args=(killswitch, c:=teraz()))
                            t.start()
                        if czy_docx:
                            img = Image.fromarray(zrzut, 'RGB')
                            img.save('temp_notuj.png')
                            dokument.add_picture('temp_notuj.png', width=Cm(16))
                            dokument.add_paragraph(scrape.strip())
                            t2.start()
                        plik.write(scrape)
                        window['stat'].update('Zanotowano... *'+str(round(zmiana, 2)))

    

    
    
    #zostaw - zamykanie
    if event == "Wyjdź" or event == sg.WIN_CLOSED:
        break


# cleanup #



window.close()

exit()



    





img = Image.fromarray(zrzut, 'RGB')
img.save('temp_notuj.png')
img.show()


#save obrazki

#get audio i analyse? / compare with ocr

#save to txt or pdf?

#controls
#set canvas

#loop it

#ew. save

main()


